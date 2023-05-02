# Functions for handling each type of files before being sent into mongodb

import os, cv2, sys
import numpy as np
from skimage.transform import rescale
from skimage.io import imread, imsave
from skimage.util import img_as_ubyte
from docx import Document

def image_encode(path:str, scale_percent:float = 0.5, desired_size:int = 300):
    if scale_percent > 1.0 and scale_percent < 0.0:
        raise Exception("scale_percent must be in range of 0.0 to 1.0!")
    if desired_size <= 0:
        raise Exception("desired size can't have the size less than 1!")
    img = imread(path)
    #if you just want to scale without border, use this instead.. will implement next year
    #resized = rescale(img,scale_percent,channel_axis=2,anti_aliasing=True) 

    # desired size with padding for ml
    # jdhao, 2017, https://jdhao.github.io/2017/11/06/resize-image-to-square-with-padding/
    old_size = img.shape[:2] # old_size is in (height, width) format

    ratio = float(desired_size)/max(old_size)
    new_size = tuple([int(x*ratio) for x in old_size])

    # new_size should be in (width, height) format
    im = cv2.resize(img, (new_size[1], new_size[0]))

    delta_w = desired_size - new_size[1]
    delta_h = desired_size - new_size[0]
    top, bottom = delta_h//2, delta_h-(delta_h//2)
    left, right = delta_w//2, delta_w-(delta_w//2)

    color = [0, 0, 0] # rpg format
    resized = cv2.copyMakeBorder(im, top, bottom, left, right, cv2.BORDER_CONSTANT,value=color)

    resized = img_as_ubyte(resized) # using cv2 color format
    resized_encode = cv2.imencode(".jpg",resized)[1] # get ndarray only
    resized_bytes = resized_encode.tobytes() # to bytes

    return resized_bytes

def image_decode(data):
    resized_decode = np.frombuffer(data,dtype=np.uint8) # create ndarray for decoding the image
    output = cv2.imdecode(resized_decode, cv2.IMREAD_COLOR) # get rescaled image
    return output

def image_write(path:str,data):
    imsave(path,data)

def docx_encode(path:str):
    with open(path,'rb') as f:
        data = Document(f)
    
    doc_str = ""
    for x in data.paragraphs:
        #print(x.text)
        doc_str += x.text + " "
    
    return doc_str

def docx_decode(path:str,data):
    doc = Document()
    doc.add_paragraph(data)
    doc.save(path)